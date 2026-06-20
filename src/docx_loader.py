from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO

from docx import Document


class DocumentLoadError(RuntimeError):
    """Raised when a DOCX document cannot be parsed."""


@dataclass(frozen=True)
class ParagraphData:
    text: str
    style_name: str
    index: int


@dataclass(frozen=True)
class DocumentData:
    file_name: str
    paragraphs: list[ParagraphData]
    full_text: str
    table_texts: list[str]
    table_count: int
    inline_shape_count: int


def load_docx(file_bytes: bytes, file_name: str) -> DocumentData:
    if not file_bytes:
        raise DocumentLoadError("파일 내용이 비어 있습니다.")

    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise DocumentLoadError(
            "유효한 DOCX 파일인지, 파일이 암호화되거나 손상되지 않았는지 확인하세요."
        ) from exc

    paragraphs = [
        ParagraphData(
            text=paragraph.text.strip(),
            style_name=paragraph.style.name if paragraph.style else "",
            index=index,
        )
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip()
    ]

    table_texts: list[str] = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            rows.append(" | ".join(cells))
        table_texts.append("\n".join(rows))

    body = "\n".join(paragraph.text for paragraph in paragraphs)
    if table_texts:
        body = f"{body}\n\n[문서 내 표]\n" + "\n\n".join(table_texts)

    return DocumentData(
        file_name=file_name,
        paragraphs=paragraphs,
        full_text=body,
        table_texts=table_texts,
        table_count=len(document.tables),
        inline_shape_count=len(document.inline_shapes),
    )
