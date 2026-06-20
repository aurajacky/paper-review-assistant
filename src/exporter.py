from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


class ExportError(RuntimeError):
    """Raised when report files cannot be created or saved."""


@dataclass(frozen=True)
class ExportedFiles:
    report_path: Path
    plan_path: Path
    report_bytes: bytes
    plan_bytes: bytes


def _safe_stem(file_name: str) -> str:
    stem = Path(file_name).stem
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", stem).strip(" .")
    return cleaned or "paper"


def _add_markdown_like_text(document: Document, text: str) -> None:
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)


def _build_report(
    file_name: str,
    created_at: datetime,
    detected_sections: list[str],
    results: dict[str, Any],
) -> bytes:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"].font.size = Pt(10)

    document.add_heading("KCI 논문 초안 점검 보고서", level=0)
    document.add_paragraph(f"분석 대상 파일명: {file_name}")
    document.add_paragraph(f"생성 시각: {created_at:%Y-%m-%d %H:%M:%S}")
    document.add_heading("감지된 논문 섹션", level=1)
    document.add_paragraph(", ".join(detected_sections) or "감지된 표준 섹션 없음")

    from src.reviewers import ANALYSIS_OPTIONS

    for key, title in ANALYSIS_OPTIONS.items():
        document.add_heading(f"{title} 점검 결과", level=1)
        analysis = results["analyses"].get(key)
        if analysis:
            _add_markdown_like_text(document, analysis.get("markdown", ""))
        else:
            document.add_paragraph("선택하지 않은 분석 항목입니다.")

    document.add_heading("수정 우선순위 Top 5", level=1)
    top_five = results.get("issues", [])[:5]
    if top_five:
        for index, issue in enumerate(top_five, start=1):
            document.add_paragraph(
                f"{index}. [{issue['severity']}] {issue['section']} - {issue['problem']}",
                style="List Number",
            )
    else:
        document.add_paragraph("구조화된 수정 이슈가 생성되지 않았습니다.")

    document.add_heading("종합 평가", level=1)
    _add_markdown_like_text(document, results.get("overall_assessment", ""))
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph(
        "주의: 이 보고서는 AI 기반 보조 점검 결과이며, 실제 투고 전에는 "
        "연구자와 지도교수의 검토가 필요합니다."
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_revision_plan(file_name: str, issues: list[dict[str, Any]]) -> bytes:
    columns = [
        "file_name",
        "section",
        "issue_type",
        "severity",
        "problem",
        "recommendation",
        "suggested_revision",
        "priority",
    ]
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "revision_plan"
    sheet.append(columns)

    for issue in issues:
        sheet.append([issue.get(column, file_name if column == "file_name" else "") for column in columns])

    header_fill = PatternFill("solid", fgColor="1F4E78")
    for cell in sheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")

    widths = [24, 20, 22, 12, 48, 48, 48, 10]
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def export_results(
    file_name: str,
    detected_sections: list[str],
    results: dict[str, Any],
    output_root: Path,
) -> ExportedFiles:
    created_at = datetime.now()
    timestamp = created_at.strftime("%Y%m%d_%H%M%S")
    stem = _safe_stem(file_name)
    report_dir = output_root / "review_reports"
    table_dir = output_root / "tables"
    report_path = report_dir / f"{stem}_{timestamp}_review_report.docx"
    plan_path = table_dir / f"{stem}_{timestamp}_revision_plan.xlsx"

    try:
        report_dir.mkdir(parents=True, exist_ok=True)
        table_dir.mkdir(parents=True, exist_ok=True)
        report_bytes = _build_report(file_name, created_at, detected_sections, results)
        plan_bytes = _build_revision_plan(file_name, results.get("issues", []))
        report_path.write_bytes(report_bytes)
        plan_path.write_bytes(plan_bytes)
    except Exception as exc:
        raise ExportError(str(exc)) from exc

    return ExportedFiles(
        report_path=report_path,
        plan_path=plan_path,
        report_bytes=report_bytes,
        plan_bytes=plan_bytes,
    )
